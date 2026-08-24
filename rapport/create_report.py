from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Rapport_Intern_Webportal_HA.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TOTAL_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("Intern webportal - HA-demo")
    set_font(r, size=9, color="59636E")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Side ")
    set_font(r, size=9, color="59636E")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT")
    set_font(r, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Intern webportal med\nhøj tilgængelighed")
    set_font(r, size=28, color="0B2545", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Linux Servere - implementering, test og evaluering")
    set_font(r, size=14, color="59636E")

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2700, 6660])
    values = [
        ("Projekt", "Intern webportal - HA-demo"),
        ("Platform", "To Proxmox VE-værter med Debian LXC-containere"),
        ("Dato", "24. august 2026"),
        ("Status", "Krav om web-HA og datareplikering er opfyldt"),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(label), size=10.5, color="0B2545", bold=True)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Afgrænsning: Backup/PBS og tredje server indgår ikke i denne rapport.")
    set_font(r, size=10, color="7A5A00", italic=True)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text), size=11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_font(p.add_run(text), size=10, color="0B2545", bold=True)
    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            p = cell.paragraphs[0]
            set_font(p.add_run(str(value)), size=9.6)
    return table


def add_report(doc):
    add_heading(doc, "1. Resumé")
    doc.add_paragraph(
        "Projektet har etableret en intern webportal på to fysiske Linux-baserede "
        "Proxmox VE-værter. Portalen er et mock-tidsregistreringssystem, der bruges "
        "til at demonstrere høj tilgængelighed, belastningsfordeling og datareplikering."
    )
    doc.add_paragraph(
        "Løsningen opfylder de centrale krav: Trafik går via én virtuel IP-adresse, "
        "fordeles mellem to webservere og fortsætter ved fejl på en proxy eller webnode. "
        "Data gemmes i PostgreSQL på en primær database og replikeres løbende til en "
        "standby-database."
    )
    add_heading(doc, "Resultat i kort form", 2)
    add_bullet(doc, "To Proxmox VE-værter er samlet i clusteret portal-ha.")
    add_bullet(doc, "VIP 192.168.1.40 flytter automatisk mellem to HAProxy/Keepalived-noder.")
    add_bullet(doc, "To Flask/Gunicorn-webnoder leverer samme mock-portal.")
    add_bullet(doc, "PostgreSQL 17 streaming-replikering er verificeret med testdata.")
    add_bullet(doc, "Test viser proxy-failover, web-failover, trafikfordeling og replikerede data.")

    add_heading(doc, "2. Opgavens mål og afgrænsning")
    doc.add_paragraph(
        "Opgaven kræver en intern webportal leveret af to Linux-servere. Trafikken skal "
        "fordeles automatisk, løsningen skal kunne overleve fejl på én server, og data "
        "oprettet via den ene webserver skal kunne ses via den anden."
    )
    add_heading(doc, "Afgrænsning", 2)
    doc.add_paragraph(
        "Portalen er bevidst et mock-system. Den kan oprette og vise testregistreringer, "
        "men indeholder ikke login, løn, personadministration eller integrationer. En "
        "eventuel tredje server til Proxmox Backup Server eller QDevice er en senere "
        "ekstraopgave og er ikke en del af denne rapport eller vurderingen her.")

    add_heading(doc, "3. Løsningsdesign")
    doc.add_paragraph(
        "De to servere er implementeret som Proxmox VE-værter. Hver vært afvikler tre "
        "LXC-containere: proxy, web og database. Løsningen er dermed opdelt i lag, så "
        "hver service kan testes og fejlfindes separat."
    )
    add_table(doc, ["Lag", "Komponenter", "Formål"], [
        ("Indgang", "portal-vip: 192.168.1.40", "Én fast adresse for brugere."),
        ("Proxy", "proxy01 / proxy02", "HAProxy fordeler HTTP-trafik; Keepalived flytter VIP'en."),
        ("Web", "web01 / web02", "Flask/Gunicorn leverer mock-tidsregistreringen."),
        ("Data", "db01 / db02", "PostgreSQL 17 med primær og fysisk streaming-standby."),
        ("Virtualisering", "pve01 / pve02", "To fysiske Proxmox VE-værter i clusteret portal-ha."),
    ], [1500, 3300, 4560])
    add_heading(doc, "Netværksplan", 2)
    add_table(doc, ["Navn", "IP-adresse", "Placering"], [
        ("pve01", "192.168.1.33", "Fysisk vært 1"),
        ("pve02", "192.168.1.34", "Fysisk vært 2"),
        ("proxy01 / proxy02", ".41 / .42", "PVE01 / PVE02"),
        ("web01 / web02", ".43 / .44", "PVE01 / PVE02"),
        ("db01 / db02", ".45 / .46", "PVE01 / PVE02"),
    ], [2700, 2400, 4260])

    add_heading(doc, "4. Implementering")
    add_heading(doc, "Platform og adgang", 2)
    doc.add_paragraph(
        "Begge værter kører samme Proxmox VE-version og er medlemmer af clusteret "
        "portal-ha. SSH-adgang er konfigureret med nøglebaseret login. Alle seks LXC-"
        "containere er sat til autostart med fast opstartsorden: database, web og proxy."
    )
    add_heading(doc, "Web og proxy", 2)
    doc.add_paragraph(
        "HAProxy kontrollerer løbende webnodernes /health-endpoint og fordeler trafik "
        "med round robin. Keepalived bruger unicast VRRP mellem proxyerne og giver den "
        "aktive proxy den virtuelle IP-adresse. Flask-applikationen viser, hvilken "
        "webserver der svarer, så belastningsfordeling kan ses direkte i browseren."
    )
    add_heading(doc, "Datareplikering", 2)
    doc.add_paragraph(
        "db01 er PostgreSQL-primær, og db02 er fysisk standby. Standbyen modtager WAL-"
        "data løbende fra primæren. Portalens endpoint /replication viser antallet af "
        "streamende standbyer; ved test rapporterede begge webnoder værdien 1."
    )

    add_heading(doc, "5. Test og kravopfyldelse")
    doc.add_paragraph("Testene er udført gennem VIP'en, så de dækker den rigtige brugervej.")
    add_table(doc, ["Krav", "Test / bevis", "Status"], [
        ("To Linux-servere leverer portal", "web01 og web02 kører på hver sin PVE-vært og svarer via VIP.", "Opfyldt"),
        ("Automatisk trafikfordeling", "To efterfølgende /health-kald ramte web01 og web02.", "Opfyldt"),
        ("Overtagelse ved proxyfejl", "Stop af Keepalived på proxy01 flyttede VIP'en til proxy02; portalen svarede fortsat.", "Opfyldt"),
        ("Overtagelse ved webfejl", "Efter HAProxy-health-check leverede web02 portalen, mens web01 var stoppet.", "Opfyldt"),
        ("Samme data via begge webservere", "Testregistrering oprettet via VIP var til stede på db02-standbyen.", "Opfyldt"),
        ("Synlig replikering", "To /replication-kald via VIP viste streaming_replicas: 1.", "Opfyldt"),
    ], [2100, 5600, 1660])
    add_heading(doc, "Observation om failover-tid", 2)
    doc.add_paragraph(
        "I den første web-failover-test blev der testet efter tre sekunder, hvor HAProxy "
        "endnu ikke havde markeret web01 som nede. Det gav HTTP 503. Efter ti sekunder "
        "leverede web02 korrekt svaret. Observationen dokumenterer, at health-check-"
        "intervallet har betydning for den oplevede failover-tid."
    )

    add_heading(doc, "6. Vurdering")
    doc.add_paragraph(
        "Opgavens centrale mål er nået. Systemet har ingen enkeltstående webserver eller "
        "proxy som ubearbejdet single point of failure, og datareplikering er både "
        "implementeret og synlig i applikationen. Mock-systemet er lille nok til at være "
        "overskueligt, men viser stadig den fulde tekniske kæde fra bruger til replikeret data."
    )
    add_heading(doc, "Bevidste begrænsninger", 2)
    add_bullet(doc, "Automatisk databasefailover er ikke aktiveret. Det kræver et sikkert konsensus- eller witness-design for at undgå split brain.")
    add_bullet(doc, "Automatisk overtagelse efter tab af en hel PVE-vært er ikke testet. Et to-noders cluster bør have en tredje QDevice-stemme, før den test udføres.")
    add_bullet(doc, "Backup, PBS og restore-test er uden for scope i denne rapport.")
    add_heading(doc, "7. Konklusion")
    doc.add_paragraph(
        "Den interne webportal er implementeret og dokumenteret som en fungerende HA-demo. "
        "Portaltrafik fordeles mellem to Linux-webnoder, VIP'en flytter ved proxyfejl, og "
        "testdata replikeres fra PostgreSQL-primæren til standbyen. De oprindelige "
        "funktionskrav er dermed nået inden for den valgte afgrænsning."
    )
    add_heading(doc, "Bilag: Dokumentationsgrundlag", 2)
    doc.add_paragraph(
        "Rapporten bygger på projektets versionsstyrede dokumentation: PROJEKTDEFINITION.md, "
        "AENDRINGSLOG.md, DRIFTSLOG.md og TESTPLAN.md samt den aktuelle implementering i GitHub-"
        "repositoryet linux-intern-webportal-ha-demo."
    )


def main():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_report(doc)
    doc.core_properties.title = "Rapport - Intern webportal med høj tilgængelighed"
    doc.core_properties.subject = "Linux Servere - HA-demo"
    doc.core_properties.author = "Projektgruppen"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
