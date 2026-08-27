"""Build the two short handover documents for the HA portal project."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).parent
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "59636E"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TOTAL_WIDTH = 9360


def font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    width_node = table._tbl.tblPr.first_child_found_in("w:tblW")
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")
    indent = table._tbl.tblPr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table._tbl.tblPr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell._tc.tcPr.tcW.set(qn("w:w"), str(width))
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(doc, preset="compact"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6 if preset == "compact" else 6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.10
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 18 if preset == "compact" else 16, 10 if preset == "compact" else 8), ("Heading 2", 13, BLUE, 14 if preset == "compact" else 12, 7 if preset == "compact" else 6), ("Heading 3", 12, DARK_BLUE, 10 if preset == "compact" else 8, 5 if preset == "compact" else 4)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "compact" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "compact" else 8)
        style.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.167
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    font(header.add_run("Intern webportal - HA-demo"), 9, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("Side "), 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def title(doc, kicker, heading, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(kicker.upper()), 10, BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(heading), 25, NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run(subtitle), 12, MUTED)


def heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    font(p.add_run(text), 11)
    return p


def add_table(doc, headers, rows, widths, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, fill)
        font(cell.paragraphs[0].add_run(value), 10, NAVY, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            font(cell.paragraphs[0].add_run(value), 9.6)
    table_geometry(table, widths)
    return table


def create_guide():
    doc = Document()
    configure(doc, "compact")
    title(doc, "Kort guide til vejledning", "Hvad er dette projekt?", "Intern webportal med høj tilgængelighed - en Linux HA-labdemonstration")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    font(p.add_run("Kort fortalt: "), 11, NAVY, bold=True)
    font(p.add_run("Jeg har bygget en lille intern tidsregistreringsportal for at bevise, at en Linux-baseret webtjeneste fortsætter, selv når en del af infrastrukturen fejler. Portalen er med vilje et mock-system: fokus er drift, stabilitet, data og dokumentation - ikke løn eller brugeradministration."), 11)
    heading(doc, "Det kan løsningen", 1)
    bullet(doc, "Fordele trafik mellem to webservere gennem én fast virtuel IP-adresse.")
    bullet(doc, "Vise både hvilken webserver der svarer, og at data er replikeret til en standby-database.")
    bullet(doc, "Overtage automatisk, hvis den aktive proxy, en webserver eller den aktuelle databaseleder fejler.")
    bullet(doc, "Bevare cluster-quorum ved tab af én fysisk Proxmox-vært.")
    bullet(doc, "Tage og gendanne både Proxmox-backup og databasekonsistent backup med WAL-historik.")
    bullet(doc, "Opdatere systemerne rullende, så redundante partnere ikke opdateres samtidigt.")
    heading(doc, "Sådan hænger det sammen", 1)
    add_table(doc, ["Lag", "Hvad det gør"], [
        ("Adgang", "Brugeren går til én virtuel IP. Keepalived flytter den til den raske proxy ved fejl."),
        ("Web", "HAProxy fordeler trafikken til web01 eller web02, som begge kører samme Flask/Gunicorn-portal."),
        ("Database", "PostgreSQL replikerer data. Patroni og tre etcd-noder vælger kun én sikker databaseleder."),
        ("Platform", "Tre Proxmox VE-værter danner clusteret portal-ha og bevarer quorum, hvis én vært forsvinder."),
        ("Backup", "PBS01 og pgBackRest/WAL giver testede gendannelsesveje med begrænset retention i labbet."),
    ], [2700, 6660])
    heading(doc, "Hvad er dokumenteret?", 1)
    doc.add_paragraph("Alle væsentlige hændelser er testet i den rigtige brugervej gennem portalens VIP. Testene viser lastfordeling, proxy- og webfailover, datareplikering, automatisk databasefailover, tab af én fysisk vært, backup/restore og rullende sikkerhedsopdatering. Efter afsluttende vedligeholdelse havde alle værter, containere og PBS01 nul ventende pakkeopdateringer.")
    heading(doc, "Vigtige afgrænsninger", 1)
    bullet(doc, "Det er et lab og et mock-system - ikke et produktionsklart tidsregistreringssystem.")
    bullet(doc, "LXC-diske ligger på lokal LVM. En container flyttes derfor ikke automatisk til en anden PVE-vært; den redundante partner fortsætter i stedet.")
    bullet(doc, "PBS01 er på den tredje labmaskine og er ikke en off-site-backup. Produktion kræver bl.a. off-site-kopi, netværkssegmentering og TLS på etcd.")
    heading(doc, "Mini-glossar", 1)
    add_table(doc, ["Begreb", "Forklaring"], [
        ("Høj tilgængelighed (HA)", "At en tjeneste fortsætter eller hurtigt overtages, når en komponent fejler."),
        ("Failover", "Automatisk overtagelse fra en fejlramt komponent til en rask reserve."),
        ("VIP", "Virtual IP: den fælles adresse brugeren benytter, selv om den aktive proxy kan skifte."),
        ("Load balancing", "Fordeling af forespørgsler mellem to webservere."),
        ("Replikering", "Løbende kopiering af databaseændringer fra leder til standby."),
        ("Quorum", "Flertal af clusterstemmer. Det forhindrer, at to dele af et cluster tror, de begge er aktive."),
        ("Split brain", "En farlig fejltilstand, hvor to noder kan skrive som leder samtidigt. Tre etcd-stemmer reducerer denne risiko."),
        ("WAL", "PostgreSQLs ændringslog, som bruges sammen med backup til konsistent gendannelse."),
    ], [2700, 6660], LIGHT_GRAY)
    doc.core_properties.title = "Hvad er dette projekt? - Intern webportal med HA"
    doc.core_properties.subject = "Kort vejlederguide og glossar"
    doc.core_properties.author = "Projektgruppen"
    out = OUT_DIR / "Kort_guide_Intern_Webportal_HA.docx"
    doc.save(out)


def create_portfolio():
    doc = Document()
    configure(doc, "business")
    title(doc, "Praktikportfolio", "Jeg byggede en webportal, der kan tåle fejl", "Linux-infrastruktur, automatisering og dokumenteret drift i et fysisk HA-lab")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(13)
    font(p.add_run("Projektet i én sætning: "), 11, NAVY, bold=True)
    font(p.add_run("Jeg designede, byggede, testede og dokumenterede en intern webportal på tre fysiske Proxmox-værter, så webtrafik og data fortsætter ved fejl på centrale komponenter."), 11)
    heading(doc, "Det har jeg selv arbejdet med", 1)
    bullet(doc, "Planlægning af arkitektur, IP-plan, scope, ændringslog og testkriterier.")
    bullet(doc, "Linux-serverdrift på Proxmox VE med Debian LXC-containere og en PBS-VM.")
    bullet(doc, "HAProxy, Keepalived og virtuel IP for load balancing og proxy-failover.")
    bullet(doc, "Flask/Gunicorn-applikation og PostgreSQL 17 med Patroni, etcd og streaming-replikering.")
    bullet(doc, "Backup/restore med Proxmox Backup Server og pgBackRest/WAL samt dokumenterede restoretests.")
    bullet(doc, "SSH-nøgler, SOPS + age til krypterede secrets samt ansvarlig opdateringspolitik.")
    heading(doc, "Det beviste jeg med test", 1)
    add_table(doc, ["Situation", "Resultat"], [
        ("Proxy eller webserver stopper", "VIP og HAProxy sender fortsat brugeren til en rask partner."),
        ("Databaseleder stopper", "Patroni promoverer sikkert replikaen, og testdata er stadig tilgængelige."),
        ("En fysisk Proxmox-vært slukkes", "De to resterende clusterstemmer bevarer quorum; portal og database fortsætter."),
        ("Data eller container skal gendannes", "Både PBS-restore og isoleret pgBackRest/WAL-restore er gennemført med kendt testdata."),
        ("Sikkerhedsopdateringer", "Udført rullende uden at tage redundant infrastruktur ned samtidigt; afsluttet med grøn status."),
    ], [3000, 6360])
    heading(doc, "Hvad det siger om mig", 1)
    doc.add_paragraph("Jeg kan omsætte et krav til en realistisk teknisk løsning, afgrænse det, bygge det i små lag, teste det ved kontrollerede fejl og dokumentere både resultater og begrænsninger. Jeg arbejder ikke kun med at få noget til at køre - jeg tænker også i drift, backup, sikkerhed, ændringsstyring og hvordan løsningen kan forklares til både teknikere og ikke-tekniske interessenter.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run("Teknologier: "), 10.5, NAVY, bold=True)
    font(p.add_run("Proxmox VE, Debian, LXC, HAProxy, Keepalived, Flask, Gunicorn, PostgreSQL, Patroni, etcd, Proxmox Backup Server, pgBackRest, SSH, SOPS + age og Git."), 10.5)
    doc.core_properties.title = "Praktikportfolio - Linux HA-lab"
    doc.core_properties.subject = "Kort projektprofil til praktikansøgning"
    doc.core_properties.author = "Projektgruppen"
    out = OUT_DIR / "Praktikportfolio_Linux_HA_lab.docx"
    doc.save(out)


if __name__ == "__main__":
    create_guide()
    create_portfolio()
    print("Created handover documents")
